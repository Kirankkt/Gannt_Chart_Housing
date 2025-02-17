import streamlit as st
import pandas as pd
import plotly.express as px
import io
import os
from datetime import datetime
from docx import Document

# ---------------------------------------------------
# App Configuration
# ---------------------------------------------------
st.set_page_config(page_title="Construction Project Manager Dashboard", layout="wide")
st.title("Construction Project Manager Dashboard II")
st.markdown(
    "This dashboard provides an overview of the project—including task snapshots, timeline visualization, and detailed reports. Use the sidebar to filter the data."
)

# ---------------------------------------------------
# 1. Data Loading from Excel
# ---------------------------------------------------
@st.cache_data
def load_data(file_path):
    if not os.path.exists(file_path):
        st.error(f"File {file_path} not found!")
        st.stop()
    df = pd.read_excel(file_path)
    # Clean column names and convert date columns
    df.columns = df.columns.str.strip()
    df["Start Date"] = pd.to_datetime(df["Start Date"], errors="coerce")
    df["End Date"] = pd.to_datetime(df["End Date"], errors="coerce")
    # Ensure Status is a string
    df["Status"] = df["Status"].astype(str)
    return df

DATA_FILE = "construction_timeline.xlsx"
df = load_data(DATA_FILE)

# ---------------------------------------------------
# 2. Data Editing Option (Direct Editing)
# ---------------------------------------------------
st.subheader("Update Task Information")
st.markdown(
    """
    **Instructions:**
    You can update any field below. For the **Status** column, please choose one of the following:
    **Finished**, **In Progress**, or **Not Started**.
    """
)

column_config = {
    "Status": st.column_config.SelectboxColumn(
        "Status",
        options=["Finished", "In Progress", "Not Started"],
        help="Select 'Finished' for completed tasks, 'In Progress' for tasks underway, or 'Not Started' for tasks that have not begun."
    )
}

edited_df = st.data_editor(df, column_config=column_config, use_container_width=True)
edited_df["Status"] = edited_df["Status"].fillna("Not Started").replace("", "Not Started")

# ---------------------------------------------------
# 2a. Save Updates Button
# ---------------------------------------------------
if st.button("Save Updates"):
    try:
        edited_df.to_excel(DATA_FILE, index=False)
        st.success("Data successfully saved!")
        load_data.clear()  # Clear cache so next load uses updated data
    except Exception as e:
        st.error(f"Error saving data: {e}")

# ---------------------------------------------------
# 2b. Manage Columns: Add and Delete Options
# ---------------------------------------------------
st.sidebar.header("Manage Columns")

# Add New Column Form
with st.sidebar.form("add_column_form"):
    new_col_name = st.text_input("New Column Name")
    default_val = st.text_input("Default Value", value="")
    add_col_submitted = st.form_submit_button("Add Column")
    if add_col_submitted:
        if new_col_name.strip() == "":
            st.sidebar.error("Please enter a valid column name.")
        elif new_col_name in edited_df.columns:
            st.sidebar.error(f"Column '{new_col_name}' already exists.")
        else:
            # Create new column, force it to string so it can handle letters/numbers
            edited_df[new_col_name] = default_val
            edited_df[new_col_name] = edited_df[new_col_name].astype(str)
            try:
                edited_df.to_excel(DATA_FILE, index=False)
                st.sidebar.success(f"Column '{new_col_name}' added successfully!")
                load_data.clear()
                if hasattr(st, "experimental_rerun"):
                    st.experimental_rerun()
                else:
                    st.sidebar.info("Please refresh the page to see the updated columns.")
            except Exception as e:
                st.sidebar.error(f"Error adding column: {e}")

# Define default columns that should never be deleted.
default_columns = {
    "Activity", "Item", "Task", "Room", "Location", "Notes",
    "Start Date", "End Date", "Status", "Workdays"
}

# Delete Column Form: only list columns that are NOT in the default set.
with st.sidebar.form("delete_column_form"):
    # Compute additional (newly added) columns only.
    additional_columns = [col for col in edited_df.columns if col not in default_columns]
    if additional_columns:
        cols_to_delete = st.multiselect("Select Newly Added Columns to Delete", options=additional_columns)
        delete_submitted = st.form_submit_button("Delete Selected Columns")
        if delete_submitted:
            if not cols_to_delete:
                st.sidebar.warning("Please select at least one column to delete.")
            else:
                edited_df.drop(columns=cols_to_delete, inplace=True)
                try:
                    edited_df.to_excel(DATA_FILE, index=False)
                    st.sidebar.success("Selected columns deleted successfully!")
                    load_data.clear()
                    if hasattr(st, "experimental_rerun"):
                        st.experimental_rerun()
                    else:
                        st.sidebar.info("Please refresh the page to see the updated columns.")
                except Exception as e:
                    st.sidebar.error(f"Error deleting columns: {e}")
    else:
        st.sidebar.info("No additional columns available for deletion.")

# ---------------------------------------------------
# 3. Sidebar Filters & Options
# ---------------------------------------------------
st.sidebar.header("Filter Options")

def norm_unique(col):
    return sorted(set(edited_df[col].dropna().astype(str).str.lower().str.strip()))

activity_options = norm_unique("Activity")
selected_activity_norm = st.sidebar.multiselect(
    "Select Activity (leave empty for all)",
    options=activity_options,
    default=[]
)

item_options = norm_unique("Item")
selected_item_norm = st.sidebar.multiselect(
    "Select Item (leave empty for all)",
    options=item_options,
    default=[]
)

task_options = norm_unique("Task")
selected_task_norm = st.sidebar.multiselect(
    "Select Task (leave empty for all)",
    options=task_options,
    default=[]
)

room_options = norm_unique("Room")
selected_room_norm = st.sidebar.multiselect(
    "Select Room (leave empty for all)",
    options=room_options,
    default=[]
)

status_options = norm_unique("Status")
selected_statuses = st.sidebar.multiselect(
    "Select Status (leave empty for all)",
    options=status_options,
    default=[]
)

show_finished = st.sidebar.checkbox("Show Finished Tasks", value=True)
color_by_status = st.sidebar.checkbox("Color-code Gantt Chart by Activity Status", value=True)

st.sidebar.markdown("**Refine Gantt Grouping**")
group_by_room = st.sidebar.checkbox("Group by Room", value=False)
group_by_item = st.sidebar.checkbox("Group by Item", value=False)
group_by_task = st.sidebar.checkbox("Group by Task", value=False)

min_date = edited_df["Start Date"].min()
max_date = edited_df["End Date"].max()
selected_date_range = st.sidebar.date_input("Select Date Range", value=[min_date, max_date])

# ---------------------------------------------------
# 4. Filtering the DataFrame Based on User Input
# ---------------------------------------------------
df_filtered = edited_df.copy()

# Create normalized columns for filtering
for col in ["Activity", "Item", "Task", "Room", "Status"]:
    df_filtered[col + "_norm"] = df_filtered[col].astype(str).str.lower().str.strip()

if selected_activity_norm:
    df_filtered = df_filtered[df_filtered["Activity_norm"].isin(selected_activity_norm)]
if selected_item_norm:
    df_filtered = df_filtered[df_filtered["Item_norm"].isin(selected_item_norm)]
if selected_task_norm:
    df_filtered = df_filtered[df_filtered["Task_norm"].isin(selected_task_norm)]
if selected_room_norm:
    df_filtered = df_filtered[df_filtered["Room_norm"].isin(selected_room_norm)]
if selected_statuses:
    df_filtered = df_filtered[df_filtered["Status_norm"].isin(selected_statuses)]
if not show_finished:
    df_filtered = df_filtered[~df_filtered["Status_norm"].eq("finished")]

if len(selected_date_range) == 2:
    start_range = pd.to_datetime(selected_date_range[0])
    end_range = pd.to_datetime(selected_date_range[1])
    df_filtered = df_filtered[
        (df_filtered["Start Date"] >= start_range) &
        (df_filtered["End Date"] <= end_range)
    ]

# Clean up
df_filtered.drop(columns=[c for c in df_filtered.columns if c.endswith("_norm")], inplace=True)

# ---------------------------------------------------
# 5. Helper Function: Compute Aggregated Status for an Activity
# ---------------------------------------------------
def aggregated_status(group_df):
    now = pd.Timestamp(datetime.today().date())
    statuses = group_df["Status"].str.strip().str.lower()
    if "in progress" in statuses.values:
        return "In Progress"
    if all(status == "finished" for status in statuses):
        max_end = group_df["End Date"].dt.normalize().max()
        if now <= max_end:
            return "Finished On Time"
        else:
            return "Finished Late"
    if all(status == "not started" for status in statuses):
        return "Not Started"
    min_start = group_df["Start Date"].dt.normalize().min()
    if now < min_start:
        return "Not Started"
    return "In Progress"

# ---------------------------------------------------
# 6. Gantt Chart Generation (Fix applied here)
# ---------------------------------------------------
def create_gantt_chart(df_input, color_by_status=False):
    """
    Build the group-by list dynamically depending on user selections.
    Always include 'Activity'. Then conditionally add Room, Item, Task.
    """
    group_cols = ["Activity"]
    if group_by_room and "Room" in df_input.columns:
        group_cols.append("Room")
    if group_by_item and "Item" in df_input.columns:
        group_cols.append("Item")
    if group_by_task and "Task" in df_input.columns:
        group_cols.append("Task")

    # Prepare a grouped DataFrame
    agg_dict = {
        "Start Date": "min",
        "End Date": "max"
    }
    if not group_cols:
        # Safety: fallback if somehow no group columns
        return px.scatter(title="No group columns selected")

    if color_by_status:
        # We'll handle color by the computed aggregated status
        agg_df = df_input.groupby(group_cols).agg(agg_dict).reset_index()

        # For each group, figure out the aggregated status
        def compute_group_status(row):
            cond = True
            for g in group_cols:
                cond = cond & (df_input[g] == row[g])
            subset = df_input[cond]
            return aggregated_status(subset)

        agg_df["Display Status"] = agg_df.apply(compute_group_status, axis=1)

        # Build a label for the y-axis from group_cols
        if len(group_cols) == 1:
            agg_df["Group Label"] = agg_df[group_cols[0]].astype(str)
        else:
            # Safely combine columns into one string
            agg_df["Group Label"] = agg_df[group_cols].apply(
                lambda row: " | ".join(row.astype(str)), axis=1
            )

        color_discrete_map = {
            "Not Started": "lightgray",
            "In Progress": "blue",
            "Finished On Time": "green",
            "Finished Late": "orange"
        }

        fig = px.timeline(
            agg_df,
            x_start="Start Date",
            x_end="End Date",
            y="Group Label",
            color="Display Status",
            color_discrete_map=color_discrete_map,
            title="Project Timeline (Color-coded by Status)"
        )
        fig.update_layout(yaxis_title=" | ".join(group_cols))

    else:
        # If not color by status, color by the top-level grouping dimension, e.g. "Activity"
        agg_df = df_input.groupby(group_cols).agg(agg_dict).reset_index()

        # Build a label
        if len(group_cols) == 1:
            agg_df["Group Label"] = agg_df[group_cols[0]].astype(str)
        else:
            agg_df["Group Label"] = agg_df[group_cols].apply(
                lambda row: " | ".join(row.astype(str)), axis=1
            )

        fig = px.timeline(
            agg_df,
            x_start="Start Date",
            x_end="End Date",
            y="Group Label",
            color=group_cols[0],  # color by the first group col, typically "Activity"
            hover_data=group_cols,
            title="Project Timeline"
        )
        fig.update_layout(yaxis_title=" | ".join(group_cols))

    fig.update_yaxes(autorange="reversed")
    fig.update_layout(xaxis_title="Timeline")
    return fig

gantt_fig = create_gantt_chart(df_filtered, color_by_status=color_by_status)

# ---------------------------------------------------
# 7. Overall Completion & Progress Bar Calculation
# ---------------------------------------------------
total_tasks = edited_df.shape[0]
finished_tasks = edited_df[edited_df["Status"].str.strip().str.lower() == "finished"].shape[0]
completion_percentage = (finished_tasks / total_tasks) * 100 if total_tasks > 0 else 0

# ---------------------------------------------------
# 8. Detailed Status Summary (for Detailed Summary Tab)
# ---------------------------------------------------
def get_status_category(status):
    s = status.strip().lower()
    if s == "finished":
        return "Finished"
    elif s == "in progress":
        return "In Progress"
    else:
        return "Not Declared"

edited_df["Status Category"] = edited_df["Status"].apply(get_status_category)
status_summary = edited_df.groupby("Status Category").size().reset_index(name="Count")
desired_order = ["Not Declared", "In Progress", "Finished"]
status_summary["Order"] = status_summary["Status Category"].apply(
    lambda x: desired_order.index(x) if x in desired_order else 99
)
status_summary = status_summary.sort_values("Order").drop("Order", axis=1)

# ---------------------------------------------------
# 9. Additional Dashboard Features
# ---------------------------------------------------
today = pd.Timestamp(datetime.today().date())
overdue_df = df_filtered[
    (df_filtered["End Date"] < today) &
    (df_filtered["Status"].str.strip().str.lower() != "finished")
]
overdue_count = overdue_df.shape[0]

task_distribution = df_filtered.groupby("Activity").size().reset_index(name="Task Count")
dist_fig = px.bar(task_distribution, x="Activity", y="Task Count", title="Task Distribution by Activity")

upcoming_start = today
upcoming_end = today + pd.Timedelta(days=7)
upcoming_df = df_filtered[
    (df_filtered["Start Date"] >= upcoming_start) &
    (df_filtered["Start Date"] <= upcoming_end)
]

filter_summary = []
if selected_activity_norm:
    filter_summary.append("Activities: " + ", ".join([s.title() for s in selected_activity_norm]))
if selected_item_norm:
    filter_summary.append("Items: " + ", ".join([s.title() for s in selected_item_norm]))
if selected_task_norm:
    filter_summary.append("Tasks: " + ", ".join([s.title() for s in selected_task_norm]))
if selected_room_norm:
    filter_summary.append("Rooms: " + ", ".join([s.title() for s in selected_room_norm]))
if selected_statuses:
    filter_summary.append("Status: " + ", ".join(selected_statuses))
if selected_date_range:
    filter_summary.append(f"Date Range: {selected_date_range[0]} to {selected_date_range[1]}")
filter_summary_text = "; ".join(filter_summary) if filter_summary else "No filters applied."

tasks_in_progress = edited_df[edited_df["Status"].str.strip().str.lower() == "in progress"].shape[0]
not_declared = edited_df[~edited_df["Status"].str.strip().str.lower().isin(["finished", "in progress"])].shape[0]

notes_df = df_filtered[df_filtered["Notes"].notna() & (df_filtered["Notes"].str.strip() != "")]

# ---------------------------------------------------
# 10. Layout with Tabs: Dashboard, Detailed Summary, Reports
# ---------------------------------------------------
tabs = st.tabs(["Dashboard", "Detailed Summary", "Reports"])

# ---------- Dashboard Tab ----------
with tabs[0]:
    st.header("Dashboard Overview")

    # 1) Show the snapshot
    st.subheader("Current Tasks Snapshot")
    st.dataframe(df_filtered)

    # 2) Show the Gantt chart below the snapshot
    st.subheader("Project Timeline")
    st.plotly_chart(gantt_fig, use_container_width=True)

    # 3) KPI & Progress
    st.metric("Overall Completion", f"{completion_percentage:.1f}%")
    st.progress(completion_percentage / 100)

    # 4) Additional Insights
    st.markdown("#### Additional Insights")
    st.markdown(f"**Overdue Tasks:** {overdue_count}")
    if not overdue_df.empty:
        st.dataframe(overdue_df[["Activity", "Room", "Task", "Status", "End Date"]])

    st.markdown("**Task Distribution by Activity:**")
    st.plotly_chart(dist_fig, use_container_width=True)

    st.markdown("**Upcoming Tasks (Next 7 Days):**")
    if not upcoming_df.empty:
        st.dataframe(upcoming_df[["Activity", "Room", "Task", "Start Date", "Status"]])
    else:
        st.info("No upcoming tasks in the next 7 days.")

    st.markdown("**Active Filters:**")
    st.write(filter_summary_text)

    col_kpi1, col_kpi2, col_kpi3, col_kpi4 = st.columns(4)
    col_kpi1.metric("Total Tasks", total_tasks)
    col_kpi2.metric("In Progress", tasks_in_progress)
    col_kpi3.metric("Finished", finished_tasks)
    col_kpi4.metric("Not Declared", not_declared)

    st.markdown("**Task Comments/Notes:**")
    if not notes_df.empty:
        st.dataframe(notes_df[["Activity", "Room", "Task", "Notes"]])
    else:
        st.info("No additional comments or notes.")

    st.markdown("Use the filters on the sidebar to adjust the view.")

# ---------- Detailed Summary Tab ----------
with tabs[1]:
    st.header("Detailed Summary")
    st.markdown("### Task Progress Tracker")
    st.dataframe(status_summary, use_container_width=True)

    st.markdown("### Full Detailed Data")
    st.dataframe(df_filtered)

# ---------- Reports Tab ----------
with tabs[2]:
    st.header("Reports")
    # Construction Daily Report Section
    st.markdown("### Construction Daily Report")

    def generate_daily_report(df):
        document = Document()
        document.add_heading("Construction Daily Report", 0)
        document.add_paragraph(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        document.add_heading("Daily Tasks", level=1)

        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Activity"
        hdr_cells[1].text = "Room"
        hdr_cells[2].text = "Task"
        hdr_cells[3].text = "Status"
        hdr_cells[4].text = "Start Date"
        hdr_cells[5].text = "End Date"

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["Activity"])
            row_cells[1].text = str(row["Room"])
            row_cells[2].text = str(row["Task"])
            row_cells[3].text = str(row["Status"])
            row_cells[4].text = (
                row["Start Date"].strftime("%Y-%m-%d") if pd.notnull(row["Start Date"]) else ""
            )
            row_cells[5].text = (
                row["End Date"].strftime("%Y-%m-%d") if pd.notnull(row["End Date"]) else ""
            )

        f = io.BytesIO()
        document.save(f)
        return f.getvalue()

    daily_report = generate_daily_report(df_filtered)
    st.download_button(
        label="Download Daily Report as Word Document",
        data=daily_report,
        file_name="Construction_Daily_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.markdown("---")

    # Change Order Template Section
    st.markdown("### Change Order Template")
    with st.form("change_order_form"):
        change_order_number = st.text_input("Change Order Number")
        project_name = st.text_input("Project Name")
        requested_by = st.text_input("Requested By")
        date = st.date_input("Date", value=datetime.today())
        change_description = st.text_area("Change Description")
        reason_for_change = st.text_area("Reason for Change")
        estimated_cost_impact = st.text_input("Estimated Cost Impact")
        approval = st.text_input("Approval (Enter approver's name)")
        submitted = st.form_submit_button("Generate Change Order Document")
        if submitted:
            form_data = {
                "change_order_number": change_order_number,
                "project_name": project_name,
                "requested_by": requested_by,
                "date": date.strftime("%Y-%m-%d"),
                "change_description": change_description,
                "reason_for_change": reason_for_change,
                "estimated_cost_impact": estimated_cost_impact,
                "approval": approval
            }
            doc = Document()
            doc.add_heading("Change Order", 0)
            doc.add_paragraph(f"Change Order Number: {form_data['change_order_number']}")
            doc.add_paragraph(f"Project Name: {form_data['project_name']}")
            doc.add_paragraph(f"Requested By: {form_data['requested_by']}")
            doc.add_paragraph(f"Date: {form_data['date']}")
            doc.add_paragraph("Change Description:")
            doc.add_paragraph(form_data['change_description'])
            doc.add_paragraph("Reason for Change:")
            doc.add_paragraph(form_data['reason_for_change'])
            doc.add_paragraph(f"Estimated Cost Impact: {form_data['estimated_cost_impact']}")
            doc.add_paragraph(f"Approval: {form_data['approval']}")
            f = io.BytesIO()
            doc.save(f)
            st.session_state["change_order_doc"] = f.getvalue()

    if "change_order_doc" in st.session_state:
        st.download_button(
            label="Download Change Order Document",
            data=st.session_state["change_order_doc"],
            file_name="Change_Order_Document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    st.markdown("---")

    # Additional Templates Section
    st.markdown("### Additional Templates")

    def generate_work_order_report(form_data):
        doc = Document()
        doc.add_heading("Work Order", 0)
        doc.add_paragraph(f"Work Order Number: {form_data['work_order_number']}")
        doc.add_paragraph(f"Contractor: {form_data['contractor']}")
        doc.add_paragraph("Work Description:")
        doc.add_paragraph(form_data['description'])
        doc.add_paragraph("Assigned Tasks:")
        doc.add_paragraph(form_data['tasks'])
        doc.add_paragraph(f"Due Date: {form_data['due_date']}")
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_risk_register_report(form_data):
        doc = Document()
        doc.add_heading("Risk Register", 0)
        doc.add_paragraph(f"Risk ID: {form_data['risk_id']}")
        doc.add_paragraph("Risk Description:")
        doc.add_paragraph(form_data['description'])
        doc.add_paragraph(f"Impact: {form_data['impact']}")
        doc.add_paragraph(f"Likelihood: {form_data['likelihood']}")
        doc.add_paragraph("Mitigation Plan:")
        doc.add_paragraph(form_data['mitigation'])
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_rfq_report(form_data):
        doc = Document()
        doc.add_heading("Request for Quote (RFQ)", 0)
        doc.add_paragraph(f"Quotation Number: {form_data['quotation_number']}")
        doc.add_paragraph(f"Customer ID: {form_data['customer_id']}")
        doc.add_paragraph(f"Company Name: {form_data['company_name']}")
        doc.add_paragraph("Requested Items/Services:")
        doc.add_paragraph(form_data['requested'])
        doc.add_paragraph(f"Quote Validity (days): {form_data['validity']}")
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_rfp_report(form_data):
        doc = Document()
        doc.add_heading("Request for Proposal (RFP)", 0)
        doc.add_paragraph(f"RFP Number: {form_data['rfp_number']}")
        doc.add_paragraph("Project Background:")
        doc.add_paragraph(form_data['background'])
        doc.add_paragraph("Scope of Work:")
        doc.add_paragraph(form_data['scope'])
        doc.add_paragraph(f"Timeline: {form_data['timeline']}")
        doc.add_paragraph(f"Submission Deadline: {form_data['deadline']}")
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_rfi_report(form_data):
        doc = Document()
        doc.add_heading("Request for Information (RFI)", 0)
        doc.add_paragraph(f"RFI Number: {form_data['rfi_number']}")
        doc.add_paragraph(f"Subject: {form_data['subject']}")
        doc.add_paragraph("Question:")
        doc.add_paragraph(form_data['question'])
        doc.add_paragraph(f"Requested Response Date: {form_data['response_date']}")
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_schedule_of_values_report(form_data):
        doc = Document()
        doc.add_heading("Schedule of Values", 0)
        doc.add_paragraph(f"Project Name: {form_data['project']}")
        doc.add_paragraph(f"Total Contract Amount: {form_data['total_amount']}")
        doc.add_paragraph("Breakdown:")
        doc.add_paragraph(form_data['breakdown'])
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_contractor_estimate_report(form_data):
        doc = Document()
        doc.add_heading("Contractor Estimate", 0)
        doc.add_paragraph(f"Estimate Number: {form_data['estimate_number']}")
        doc.add_paragraph(f"Project Name: {form_data['project']}")
        doc.add_paragraph(f"Estimated Material Costs: {form_data['material_costs']}")
        doc.add_paragraph(f"Estimated Labor Costs: {form_data['labor_costs']}")
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_construction_quote_report(form_data):
        doc = Document()
        doc.add_heading("Construction Quote", 0)
        doc.add_paragraph(f"Quote Number: {form_data['quote_number']}")
        doc.add_paragraph(f"Project Name: {form_data['project']}")
        doc.add_paragraph(f"Estimated Total Cost: {form_data['total_cost']}")
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_scope_of_work_report(form_data):
        doc = Document()
        doc.add_heading("Scope of Work", 0)
        doc.add_paragraph(f"Project Name: {form_data['project_name']}")
        doc.add_paragraph("Scope Details:")
        doc.add_paragraph(form_data['scope_details'])
        doc.add_paragraph("Milestones and Deliverables:")
        doc.add_paragraph(form_data['milestones'])
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_painting_estimate_report(form_data):
        doc = Document()
        doc.add_heading("Painting Estimate", 0)
        doc.add_paragraph(f"Estimate Number: {form_data['estimate_number']}")
        doc.add_paragraph(f"Project/Location: {form_data['project']}")
        doc.add_paragraph(f"Estimated Material Costs: {form_data['material_costs']}")
        doc.add_paragraph(f"Estimated Labor Costs: {form_data['labor_costs']}")
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def generate_roofing_estimate_report(form_data):
        doc = Document()
        doc.add_heading("Roofing Estimate", 0)
        doc.add_paragraph(f"Estimate Number: {form_data['estimate_number']}")
        doc.add_paragraph(f"Total Area (sq ft): {form_data['area']}")
        doc.add_paragraph(f"Material Specification: {form_data['materials']}")
        doc.add_paragraph(f"Estimated Cost: {form_data['estimated_cost']}")
        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    template_choice = st.selectbox(
        "Select Template to Generate",
        options=[
            "Work Order Template",
            "Risk Register Template",
            "Request for Quote (RFQ) Template",
            "Request for Proposal (RFP) Template",
            "Request for Information (RFI) Template",
            "Schedule of Values Template",
            "Contractor Estimate Template",
            "Construction Quote Template",
            "Scope of Work Template",
            "Painting Estimate Template",
            "Roofing Estimate Template"
        ]
    )

    if template_choice == "Work Order Template":
        with st.form("work_order_form"):
            work_order_number = st.text_input("Work Order Number")
            contractor = st.text_input("Contractor")
            description = st.text_area("Work Description")
            tasks_input = st.text_area("Assigned Tasks")
            due_date = st.date_input("Due Date", value=datetime.today())
            submitted = st.form_submit_button("Generate Work Order Document")
            if submitted:
                form_data = {
                    "work_order_number": work_order_number,
                    "contractor": contractor,
                    "description": description,
                    "tasks": tasks_input,
                    "due_date": due_date.strftime("%Y-%m-%d")
                }
                doc_bytes = generate_work_order_report(form_data)
                st.session_state["work_order_doc"] = doc_bytes
        if "work_order_doc" in st.session_state:
            st.download_button(
                label="Download Work Order Document",
                data=st.session_state["work_order_doc"],
                file_name="Work_Order_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Risk Register Template":
        with st.form("risk_register_form"):
            risk_id = st.text_input("Risk ID")
            description = st.text_area("Risk Description")
            impact = st.text_input("Impact")
            likelihood = st.text_input("Likelihood")
            mitigation = st.text_area("Mitigation Plan")
            submitted = st.form_submit_button("Generate Risk Register Document")
            if submitted:
                form_data = {
                    "risk_id": risk_id,
                    "description": description,
                    "impact": impact,
                    "likelihood": likelihood,
                    "mitigation": mitigation
                }
                doc_bytes = generate_risk_register_report(form_data)
                st.session_state["risk_register_doc"] = doc_bytes
        if "risk_register_doc" in st.session_state:
            st.download_button(
                label="Download Risk Register Document",
                data=st.session_state["risk_register_doc"],
                file_name="Risk_Register_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Request for Quote (RFQ) Template":
        with st.form("rfq_form"):
            quotation_number = st.text_input("Quotation Number")
            customer_id = st.text_input("Customer ID")
            company_name = st.text_input("Company Name")
            requested = st.text_area("Requested Items/Services")
            validity = st.text_input("Quote Validity (days)")
            submitted = st.form_submit_button("Generate RFQ Document")
            if submitted:
                form_data = {
                    "quotation_number": quotation_number,
                    "customer_id": customer_id,
                    "company_name": company_name,
                    "requested": requested,
                    "validity": validity
                }
                doc_bytes = generate_rfq_report(form_data)
                st.session_state["rfq_doc"] = doc_bytes
        if "rfq_doc" in st.session_state:
            st.download_button(
                label="Download RFQ Document",
                data=st.session_state["rfq_doc"],
                file_name="RFQ_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Request for Proposal (RFP) Template":
        with st.form("rfp_form"):
            rfp_number = st.text_input("RFP Number")
            background = st.text_area("Project Background")
            scope = st.text_area("Scope of Work")
            timeline = st.text_input("Timeline")
            deadline = st.date_input("Submission Deadline", value=datetime.today())
            submitted = st.form_submit_button("Generate RFP Document")
            if submitted:
                form_data = {
                    "rfp_number": rfp_number,
                    "background": background,
                    "scope": scope,
                    "timeline": timeline,
                    "deadline": deadline.strftime("%Y-%m-%d")
                }
                doc_bytes = generate_rfp_report(form_data)
                st.session_state["rfp_doc"] = doc_bytes
        if "rfp_doc" in st.session_state:
            st.download_button(
                label="Download RFP Document",
                data=st.session_state["rfp_doc"],
                file_name="RFP_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Request for Information (RFI) Template":
        with st.form("rfi_form"):
            rfi_number = st.text_input("RFI Number")
            subject = st.text_input("Subject")
            question = st.text_area("Question")
            response_date = st.date_input("Requested Response Date", value=datetime.today())
            submitted = st.form_submit_button("Generate RFI Document")
            if submitted:
                form_data = {
                    "rfi_number": rfi_number,
                    "subject": subject,
                    "question": question,
                    "response_date": response_date.strftime("%Y-%m-%d")
                }
                doc_bytes = generate_rfi_report(form_data)
                st.session_state["rfi_doc"] = doc_bytes
        if "rfi_doc" in st.session_state:
            st.download_button(
                label="Download RFI Document",
                data=st.session_state["rfi_doc"],
                file_name="RFI_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Schedule of Values Template":
        with st.form("schedule_values_form"):
            project = st.text_input("Project Name")
            total_amount = st.text_input("Total Contract Amount")
            breakdown = st.text_area("Task Breakdown (list tasks and amounts)")
            submitted = st.form_submit_button("Generate Schedule of Values Document")
            if submitted:
                form_data = {
                    "project": project,
                    "total_amount": total_amount,
                    "breakdown": breakdown
                }
                doc_bytes = generate_schedule_of_values_report(form_data)
                st.session_state["schedule_doc"] = doc_bytes
        if "schedule_doc" in st.session_state:
            st.download_button(
                label="Download Schedule of Values Document",
                data=st.session_state["schedule_doc"],
                file_name="Schedule_of_Values_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Contractor Estimate Template":
        with st.form("contractor_estimate_form"):
            estimate_number = st.text_input("Estimate Number")
            project = st.text_input("Project Name")
            material_costs = st.text_input("Estimated Material Costs")
            labor_costs = st.text_input("Estimated Labor Costs")
            submitted = st.form_submit_button("Generate Contractor Estimate Document")
            if submitted:
                form_data = {
                    "estimate_number": estimate_number,
                    "project": project,
                    "material_costs": material_costs,
                    "labor_costs": labor_costs
                }
                doc_bytes = generate_contractor_estimate_report(form_data)
                st.session_state["contractor_estimate_doc"] = doc_bytes
        if "contractor_estimate_doc" in st.session_state:
            st.download_button(
                label="Download Contractor Estimate Document",
                data=st.session_state["contractor_estimate_doc"],
                file_name="Contractor_Estimate_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Construction Quote Template":
        with st.form("construction_quote_form"):
            quote_number = st.text_input("Quote Number")
            project = st.text_input("Project Name")
            total_cost = st.text_input("Estimated Total Cost")
            submitted = st.form_submit_button("Generate Construction Quote Document")
            if submitted:
                form_data = {
                    "quote_number": quote_number,
                    "project": project,
                    "total_cost": total_cost
                }
                doc_bytes = generate_construction_quote_report(form_data)
                st.session_state["construction_quote_doc"] = doc_bytes
        if "construction_quote_doc" in st.session_state:
            st.download_button(
                label="Download Construction Quote Document",
                data=st.session_state["construction_quote_doc"],
                file_name="Construction_Quote_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Scope of Work Template":
        with st.form("scope_work_form"):
            project_name = st.text_input("Project Name")
            scope_details = st.text_area("Scope Details")
            milestones = st.text_area("Milestones and Deliverables")
            submitted = st.form_submit_button("Generate Scope of Work Document")
            if submitted:
                form_data = {
                    "project_name": project_name,
                    "scope_details": scope_details,
                    "milestones": milestones
                }
                doc_bytes = generate_scope_of_work_report(form_data)
                st.session_state["scope_work_doc"] = doc_bytes
        if "scope_work_doc" in st.session_state:
            st.download_button(
                label="Download Scope of Work Document",
                data=st.session_state["scope_work_doc"],
                file_name="Scope_of_Work_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Painting Estimate Template":
        with st.form("painting_estimate_form"):
            estimate_number = st.text_input("Estimate Number")
            project = st.text_input("Project/Location")
            material_costs = st.text_input("Estimated Material Costs")
            labor_costs = st.text_input("Estimated Labor Costs")
            submitted = st.form_submit_button("Generate Painting Estimate Document")
            if submitted:
                form_data = {
                    "estimate_number": estimate_number,
                    "project": project,
                    "material_costs": material_costs,
                    "labor_costs": labor_costs
                }
                doc_bytes = generate_painting_estimate_report(form_data)
                st.session_state["painting_estimate_doc"] = doc_bytes
        if "painting_estimate_doc" in st.session_state:
            st.download_button(
                label="Download Painting Estimate Document",
                data=st.session_state["painting_estimate_doc"],
                file_name="Painting_Estimate_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif template_choice == "Roofing Estimate Template":
        with st.form("roofing_estimate_form"):
            estimate_number = st.text_input("Estimate Number")
            area = st.text_input("Total Area (sq ft)")
            materials = st.text_input("Material Specification")
            estimated_cost = st.text_input("Estimated Cost")
            submitted = st.form_submit_button("Generate Roofing Estimate Document")
            if submitted:
                form_data = {
                    "estimate_number": estimate_number,
                    "area": area,
                    "materials": materials,
                    "estimated_cost": estimated_cost
                }
                doc_bytes = generate_roofing_estimate_report(form_data)
                st.session_state["roofing_estimate_doc"] = doc_bytes
        if "roofing_estimate_doc" in st.session_state:
            st.download_button(
                label="Download Roofing Estimate Document",
                data=st.session_state["roofing_estimate_doc"],
                file_name="Roofing_Estimate_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    st.markdown("---")

    # Export Data
    st.markdown("### Export Data")

    def convert_df_to_csv(df):
        return df.to_csv(index=False).encode("utf-8")

    def convert_df_to_excel(df):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="FilteredData")
        return output.getvalue()

    csv_data = convert_df_to_csv(df_filtered)
    st.download_button(
        label="Download Filtered Data as CSV",
        data=csv_data,
        file_name="filtered_construction_data.csv",
        mime="text/csv"
    )

    excel_data = convert_df_to_excel(df_filtered)
    st.download_button(
        label="Download Filtered Data as Excel",
        data=excel_data,
        file_name="filtered_construction_data.xlsx",
        mime="application/vnd.ms-excel"
    )

st.markdown("---")
st.markdown("CMBP Analytics II")
